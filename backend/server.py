from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form, Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
from emergentintegrations.llm.chat import LlmChat, UserMessage
from pinecone import Pinecone, ServerlessSpec
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
import json
import shutil

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME', 'visar_immigration_db')]

# File upload directory
UPLOAD_DIR = ROOT_DIR / 'uploads'
UPLOAD_DIR.mkdir(exist_ok=True)

# JWT Secret
JWT_SECRET = os.environ.get('JWT_SECRET', 'visar-secret-key-change-in-production')
JWT_ALGORITHM = 'HS256'

# Initialize Pinecone
pc = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))
INDEX_NAME = 'visar-petitions'

# Create index if it doesn't exist
try:
    if INDEX_NAME not in pc.list_indexes().names():
        pc.create_index(
            name=INDEX_NAME,
            dimension=384,  # all-MiniLM-L6-v2 embedding dimension
            metric='cosine',
            spec=ServerlessSpec(cloud='aws', region='us-east-1')
        )
except Exception as e:
    logging.error(f"Pinecone initialization error: {e}")

# Initialize embedding model
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Security
security = HTTPBearer()

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Models
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    password_hash: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserRegister(BaseModel):
    username: str
    password: str

class UserLogin(BaseModel):
    username: str
    password: str

class TokenResponse(BaseModel):
    token: str
    username: str

class Client(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    name: str
    email: str
    visa_type: str  # EB1A, EB2NIW, O-1A
    status: str = "active"  # active, archived
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ClientCreate(BaseModel):
    name: str
    email: str
    visa_type: str

class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    filename: str
    file_path: str
    file_type: str  # cv, evidence, training
    uploaded_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Template(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    visa_type: str  # EB1A, EB2NIW, O-1A
    criterion: str  # awards, media, judging, original_contributions, etc.
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TemplateCreate(BaseModel):
    visa_type: str
    criterion: str
    content: str

class Petition(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    visa_type: str
    criterion: Optional[str] = None
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class PetitionGenerate(BaseModel):
    client_id: str
    visa_type: str
    criterion: Optional[str] = None
    prompt: str

class TrainingDoc(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_path: str
    doc_type: str  # successful, unsuccessful
    visa_type: str
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatMessage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    role: str  # user, assistant
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatRequest(BaseModel):
    client_id: str
    message: str

class ChatResponse(BaseModel):
    response: str
    message_id: str

# Helper Functions
def create_token(username: str) -> str:
    payload = {
        'username': username,
        'exp': datetime.now(timezone.utc) + timedelta(days=7)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)) -> str:
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload['username']
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

def extract_text_from_pdf(file_path: str) -> str:
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text
    except:
        return ''

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except:
        return ''

async def get_relevant_context(query: str, top_k: int = 3) -> str:
    """Get relevant context from training documents using Pinecone"""
    try:
        index = pc.Index(INDEX_NAME)
        query_embedding = embedding_model.encode(query).tolist()
        
        results = index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True
        )
        
        context_parts = []
        for match in results.matches:
            if match.score > 0.5:  # Only include relevant matches
                context_parts.append(match.metadata.get('text', ''))
        
        return '\n\n---\n\n'.join(context_parts) if context_parts else ''
    except Exception as e:
        logging.error(f"Error getting context: {e}")
        return ''

# Auth Endpoints
@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserRegister):
    # Check if user exists
    existing = await db.users.find_one({'username': user_data.username})
    if existing:
        raise HTTPException(status_code=400, detail="Username already exists")
    
    # Hash password
    password_hash = bcrypt.hashpw(user_data.password.encode(), bcrypt.gensalt()).decode()
    
    user = User(username=user_data.username, password_hash=password_hash)
    doc = user.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.users.insert_one(doc)
    
    token = create_token(user.username)
    return TokenResponse(token=token, username=user.username)

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(user_data: UserLogin):
    user = await db.users.find_one({'username': user_data.username})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not bcrypt.checkpw(user_data.password.encode(), user['password_hash'].encode()):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_token(user['username'])
    return TokenResponse(token=token, username=user['username'])

# Client Endpoints
@api_router.post("/clients", response_model=Client)
async def create_client(client_data: ClientCreate, username: str = Depends(verify_token)):
    user = await db.users.find_one({'username': username})
    
    client = Client(
        user_id=user['id'],
        name=client_data.name,
        email=client_data.email,
        visa_type=client_data.visa_type
    )
    
    doc = client.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.clients.insert_one(doc)
    return client

@api_router.get("/clients", response_model=List[Client])
async def get_clients(username: str = Depends(verify_token)):
    user = await db.users.find_one({'username': username})
    clients = await db.clients.find({'user_id': user['id']}, {'_id': 0}).to_list(1000)
    
    for client in clients:
        if isinstance(client['created_at'], str):
            client['created_at'] = datetime.fromisoformat(client['created_at'])
    
    return clients

@api_router.get("/clients/{client_id}", response_model=Client)
async def get_client(client_id: str, username: str = Depends(verify_token)):
    client = await db.clients.find_one({'id': client_id}, {'_id': 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    if isinstance(client['created_at'], str):
        client['created_at'] = datetime.fromisoformat(client['created_at'])
    
    return client

# Document Upload Endpoints
@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    client_id: str = Form(...),
    file_type: str = Form(...),
    username: str = Depends(verify_token)
):
    # Create client directory
    client_dir = UPLOAD_DIR / client_id
    client_dir.mkdir(exist_ok=True)
    
    # Save file
    file_path = client_dir / file.filename
    with open(file_path, 'wb') as f:
        shutil.copyfileobj(file.file, f)
    
    # Save metadata
    document = Document(
        client_id=client_id,
        filename=file.filename,
        file_path=str(file_path),
        file_type=file_type
    )
    
    doc = document.model_dump()
    doc['uploaded_at'] = doc['uploaded_at'].isoformat()
    
    await db.documents.insert_one(doc)
    
    return {"message": "File uploaded successfully", "document_id": document.id}

@api_router.get("/documents/{client_id}", response_model=List[Document])
async def get_documents(client_id: str, username: str = Depends(verify_token)):
    docs = await db.documents.find({'client_id': client_id}, {'_id': 0}).to_list(1000)
    
    for doc in docs:
        if isinstance(doc['uploaded_at'], str):
            doc['uploaded_at'] = datetime.fromisoformat(doc['uploaded_at'])
    
    return docs

# Template Endpoints
@api_router.post("/templates", response_model=Template)
async def create_template(template_data: TemplateCreate, username: str = Depends(verify_token)):
    template = Template(
        visa_type=template_data.visa_type,
        criterion=template_data.criterion,
        content=template_data.content
    )
    
    doc = template.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.templates.insert_one(doc)
    return template

@api_router.get("/templates", response_model=List[Template])
async def get_templates(visa_type: Optional[str] = None, username: str = Depends(verify_token)):
    query = {'visa_type': visa_type} if visa_type else {}
    templates = await db.templates.find(query, {'_id': 0}).to_list(1000)
    
    for template in templates:
        if isinstance(template['created_at'], str):
            template['created_at'] = datetime.fromisoformat(template['created_at'])
    
    return templates

# Petition Generation Endpoint
@api_router.post("/petitions/generate", response_model=ChatResponse)
async def generate_petition(request: PetitionGenerate, username: str = Depends(verify_token)):
    # Get client info
    client = await db.clients.find_one({'id': request.client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Get relevant templates
    templates = await db.templates.find({
        'visa_type': request.visa_type,
        'criterion': request.criterion
    }, {'_id': 0}).to_list(10)
    
    template_context = '\n\n'.join([t['content'] for t in templates]) if templates else ''
    
    # Get relevant context from training docs
    rag_context = await get_relevant_context(request.prompt)
    
    # Build system message
    template_section = f'Template Context:\n{template_context}\n\n' if template_context else ''
    rag_section = f'Reference Context from Successful Petitions:\n{rag_context}\n\n' if rag_context else ''
    
    system_message = f"""You are an expert immigration petition writer for Visar, specializing in {request.visa_type} visas.

Your writing style is:
- Professional and persuasive
- Evidence-backed and detailed
- Aligned with USCIS requirements
- Clear and compelling

Client: {client['name']}
Visa Type: {request.visa_type}
Criterion: {request.criterion or 'General'}

{template_section}{rag_section}
Generate a high-quality petition section based on the user's request."""
    
    # Generate with GPT-4o
    chat = LlmChat(
        api_key=os.environ.get('EMERGENT_LLM_KEY'),
        session_id=f"petition_{request.client_id}",
        system_message=system_message
    ).with_model("openai", "gpt-4o")
    
    user_message = UserMessage(text=request.prompt)
    response = await chat.send_message(user_message)
    
    # Save petition
    petition = Petition(
        client_id=request.client_id,
        visa_type=request.visa_type,
        criterion=request.criterion,
        content=response
    )
    
    doc = petition.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.petitions.insert_one(doc)
    
    return ChatResponse(response=response, message_id=petition.id)

@api_router.get("/petitions/{client_id}", response_model=List[Petition])
async def get_petitions(client_id: str, username: str = Depends(verify_token)):
    petitions = await db.petitions.find({'client_id': client_id}, {'_id': 0}).to_list(1000)
    
    for petition in petitions:
        if isinstance(petition['created_at'], str):
            petition['created_at'] = datetime.fromisoformat(petition['created_at'])
    
    return petitions

# Training Document Upload
@api_router.post("/training/upload")
async def upload_training_doc(
    file: UploadFile = File(...),
    doc_type: str = Form(...),
    visa_type: str = Form(...),
    username: str = Depends(verify_token)
):
    # Save file
    training_dir = UPLOAD_DIR / 'training'
    training_dir.mkdir(exist_ok=True)
    
    file_path = training_dir / file.filename
    with open(file_path, 'wb') as f:
        shutil.copyfileobj(file.file, f)
    
    # Extract text
    if file.filename.endswith('.pdf'):
        content = extract_text_from_pdf(str(file_path))
    elif file.filename.endswith('.docx'):
        content = extract_text_from_docx(str(file_path))
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    
    # Save to database
    training_doc = TrainingDoc(
        filename=file.filename,
        file_path=str(file_path),
        doc_type=doc_type,
        visa_type=visa_type,
        content=content
    )
    
    doc = training_doc.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.training_docs.insert_one(doc)
    
    # Create embedding and store in Pinecone
    try:
        index = pc.Index(INDEX_NAME)
        embedding = embedding_model.encode(content).tolist()
        
        index.upsert(vectors=[(
            training_doc.id,
            embedding,
            {
                'text': content[:1000],  # Store first 1000 chars
                'visa_type': visa_type,
                'doc_type': doc_type,
                'filename': file.filename
            }
        )])
    except Exception as e:
        logging.error(f"Error storing in Pinecone: {e}")
    
    return {"message": "Training document uploaded successfully", "doc_id": training_doc.id}

@api_router.get("/training/docs", response_model=List[TrainingDoc])
async def get_training_docs(username: str = Depends(verify_token)):
    docs = await db.training_docs.find({}, {'_id': 0}).to_list(1000)
    
    for doc in docs:
        if isinstance(doc['created_at'], str):
            doc['created_at'] = datetime.fromisoformat(doc['created_at'])
    
    return docs

# Chat Endpoints
@api_router.post("/chat/message", response_model=ChatResponse)
async def send_chat_message(request: ChatRequest, username: str = Depends(verify_token)):
    # Get client
    client = await db.clients.find_one({'id': request.client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Save user message
    user_msg = ChatMessage(
        client_id=request.client_id,
        role='user',
        content=request.message
    )
    user_doc = user_msg.model_dump()
    user_doc['created_at'] = user_doc['created_at'].isoformat()
    await db.chat_history.insert_one(user_doc)
    
    # Get chat history
    history = await db.chat_history.find(
        {'client_id': request.client_id},
        {'_id': 0}
    ).sort('created_at', 1).to_list(50)
    
    # Get relevant context
    rag_context = await get_relevant_context(request.message)
    
    # Build system message
    rag_section = f'Reference Context:\n{rag_context}\n\n' if rag_context else ''
    
    system_message = f"""You are an expert immigration petition assistant for Visar, helping with {client['visa_type']} visa petitions.

Client: {client['name']}
Visa Type: {client['visa_type']}

{rag_section}
Provide helpful, professional advice and draft petition content as needed."""
    
    # Generate response
    chat = LlmChat(
        api_key=os.environ.get('EMERGENT_LLM_KEY'),
        session_id=f"chat_{request.client_id}",
        system_message=system_message
    ).with_model("openai", "gpt-5")
    
    user_message = UserMessage(text=request.message)
    response = await chat.send_message(user_message)
    
    # Save assistant message
    assistant_msg = ChatMessage(
        client_id=request.client_id,
        role='assistant',
        content=response
    )
    assistant_doc = assistant_msg.model_dump()
    assistant_doc['created_at'] = assistant_doc['created_at'].isoformat()
    await db.chat_history.insert_one(assistant_doc)
    
    return ChatResponse(response=response, message_id=assistant_msg.id)

@api_router.get("/chat/history/{client_id}", response_model=List[ChatMessage])
async def get_chat_history(client_id: str, username: str = Depends(verify_token)):
    history = await db.chat_history.find(
        {'client_id': client_id},
        {'_id': 0}
    ).sort('created_at', 1).to_list(1000)
    
    for msg in history:
        if isinstance(msg['created_at'], str):
            msg['created_at'] = datetime.fromisoformat(msg['created_at'])
    
    return history

# Health check
@api_router.get("/")
async def root():
    return {"message": "Visar Immigration API"}

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
